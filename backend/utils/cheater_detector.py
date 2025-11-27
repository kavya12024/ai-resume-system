import pdfplumber
import docx2txt
import os
import re
from docx import Document
from docx.shared import RGBColor, Pt

def extract_text_from_file(file_path):
    print(f"\n🔥🔥🔥 ANALYZING: {os.path.basename(file_path)} 🔥🔥🔥")
    
    text = ""
    is_suspicious = False
    hidden_count = 0
    
    extension = os.path.splitext(file_path)[1].lower()

    try:
        # ==========================================
        # 📄 LOGIC FOR PDF FILES
        # ==========================================
        if extension == '.pdf':
            print("--- SCANNING PDF LAYER ---")
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    chars_to_keep = []
                    for char in page.chars:
                        is_hidden = False
                        
                        # 1. Size Check (< 6pt is suspicious)
                        if char.get('size', 0) < 6: is_hidden = True
                        
                        # 2. Color Check (White)
                        color = char.get('non_stroking_color')
                        if color == (1, 1, 1) or color == 1: is_hidden = True
                        if isinstance(color, tuple) and len(color) == 3:
                            if all(c > 0.95 for c in color): is_hidden = True

                        if is_hidden:
                            hidden_count += 1
                        else:
                            # Only keep visible characters
                            chars_to_keep.append(char)
                    
                    # Extract text ONLY from the clean characters
                    page_text = pdfplumber.utils.extract_text(chars_to_keep, x_tolerance=2, y_tolerance=2)
                    if page_text:
                        text += page_text + "\n"

        # ==========================================
        # 📝 LOGIC FOR DOCX FILES (UPDATED!)
        # ==========================================
        elif extension == '.docx':
            print("--- SCANNING DOCX XML ---")
            
            doc = Document(file_path)
            docx_text_parts = []

            # Helper function to check a "Run" (a chunk of text)
            def is_run_suspicious(run):
                # Check 1: Tiny Font (< 6pt)
                if run.font.size and run.font.size.pt < 6:
                    print(f"   [!] Skipped Tiny Text: '{run.text[:15]}...'")
                    return True
                
                # Check 2: White Color
                if run.font.color and run.font.color.rgb == RGBColor(255, 255, 255):
                    print(f"   [!] Skipped White Text: '{run.text[:15]}...'")
                    return True
                
                # Check 3: Hidden Attribute
                if run.font.hidden:
                    print(f"   [!] Skipped Hidden Attribute: '{run.text[:15]}...'")
                    return True
                    
                return False

            # 1. Scan Paragraphs
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if is_run_suspicious(run):
                        hidden_count += 10
                    else:
                        docx_text_parts.append(run.text)
                docx_text_parts.append("\n") # Preserve line breaks

            # 2. Scan Tables (Important! Resumes use tables)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if is_run_suspicious(run):
                                    hidden_count += 10
                                else:
                                    docx_text_parts.append(run.text)
                            docx_text_parts.append(" ") # Space between cells

            # Join the valid parts together
            text = "".join(docx_text_parts)

        else:
            return None, False
        
        # ==========================================
        # 🏁 FINAL DECISION
        # ==========================================
        print(f"--- TOTAL SUSPICIOUS SCORE: {hidden_count} ---")
        
        if hidden_count > 10:
            is_suspicious = True
            print("🚨 MALPRACTICE DETECTED: Flagged as Suspicious 🚨")
        
        # Cleanup
        text = re.sub(r'\s+', ' ', text).strip()
        
        return text, is_suspicious

    except Exception as e:
        print(f"Error extracting text: {e}")
        return None, False